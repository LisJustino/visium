from pathlib import Path
import sys
import zipfile

import fitz
from docx import Document


# =========================================================
# VISIUM
# INVENTÁRIO E EXTRAÇÃO DE IMAGENS
#
# Suporta:
# - PDF
# - DOCX
# =========================================================


ROOT_DIR = Path(__file__).resolve().parent.parent

DOCS_DIR = ROOT_DIR / "docs"

IMAGES_DIR = (
    ROOT_DIR
    / "assets"
    / "img"
)

REPORTS_DIR = (
    ROOT_DIR
    / "docs"
    / "development"
    / "reports"
)


# =========================================================
# UTILITÁRIOS
# =========================================================


def normalizar_nome(
    nome: str
) -> str:
    """
    Converte um nome para um formato seguro
    para pasta ou arquivo.
    """

    return (
        nome
        .strip()
        .lower()
        .replace(" ", "-")
        .replace("_", "-")
    )


def encontrar_documento(
    nome_documento: str
) -> Path:
    """
    Procura um PDF ou DOCX dentro da pasta docs.
    """

    caminho = (
        DOCS_DIR
        / nome_documento
    )

    if caminho.exists():

        return caminho


    extensoes = [
        ".pdf",
        ".docx"
    ]


    if not Path(
        nome_documento
    ).suffix:

        for extensao in extensoes:

            candidato = (
                DOCS_DIR
                / f"{nome_documento}{extensao}"
            )

            if candidato.exists():

                return candidato


    raise FileNotFoundError(
        "Documento não encontrado em: "
        f"{DOCS_DIR}"
    )


def formatar_tamanho(
    valor: int
) -> str:
    """
    Formata tamanho em bytes
    para leitura humana.
    """

    unidades = [
        "B",
        "KB",
        "MB",
        "GB"
    ]

    tamanho = float(
        valor
    )


    for unidade in unidades:

        if tamanho < 1024:

            return (
                f"{tamanho:.1f} "
                f"{unidade}"
            )


        tamanho /= 1024


    return (
        f"{tamanho:.1f} TB"
    )


def garantir_pasta_relatorios() -> None:
    """
    Cria a pasta de relatórios
    caso ela não exista.
    """

    REPORTS_DIR.mkdir(
        parents=True,
        exist_ok=True
    )


# =========================================================
# INVENTÁRIO — PDF
# =========================================================


def inventariar_pdf(
    pdf_path: Path
) -> str:
    """
    Analisa um PDF e gera
    um relatório textual.
    """

    documento = fitz.open(
        pdf_path
    )

    total_paginas = len(
        documento
    )

    linhas = []


    linhas.append(
        "=" * 72
    )

    linhas.append(
        "VISIUM — INVENTÁRIO DE MATERIAL"
    )

    linhas.append(
        "=" * 72
    )

    linhas.append("")

    linhas.append(
        f"Arquivo: {pdf_path.name}"
    )

    linhas.append(
        f"Tipo: PDF"
    )

    linhas.append(
        f"Páginas: {total_paginas}"
    )

    linhas.append("")

    linhas.append(
        "-" * 72
    )


    total_imagens = 0

    total_desenhos = 0

    paginas_com_texto = 0

    paginas_com_imagens = 0


    for numero_pagina in range(
        total_paginas
    ):

        pagina = documento[
            numero_pagina
        ]


        texto = pagina.get_text(
            "text"
        ).strip()


        imagens = pagina.get_images(
            full=True
        )


        desenhos = (
            pagina.get_drawings()
        )


        if texto:

            paginas_com_texto += 1


        if imagens:

            paginas_com_imagens += 1


        total_imagens += len(
            imagens
        )


        total_desenhos += len(
            desenhos
        )


        linhas.append("")

        linhas.append(
            f"PÁGINA "
            f"{numero_pagina + 1:03d}"
        )

        linhas.append(
            "-" * 40
        )


        linhas.append(
            "Texto: "
            + (
                "sim"
                if texto
                else "não"
            )
        )


        linhas.append(
            "Imagens incorporadas: "
            f"{len(imagens)}"
        )


        linhas.append(
            "Elementos vetoriais: "
            f"{len(desenhos)}"
        )


        if texto:

            texto_linhas = [
                linha.strip()
                for linha in texto.splitlines()
                if linha.strip()
            ]


            if texto_linhas:

                linhas.append(
                    "Primeiro conteúdo textual:"
                )


                limite = min(
                    5,
                    len(texto_linhas)
                )


                for linha in texto_linhas[
                    :limite
                ]:

                    linha_normalizada = (
                        linha
                        .replace(
                            "\t",
                            " "
                        )
                        .strip()
                    )


                    if len(
                        linha_normalizada
                    ) > 140:

                        linha_normalizada = (
                            linha_normalizada[:137]
                            + "..."
                        )


                    linhas.append(
                        f"  - "
                        f"{linha_normalizada}"
                    )


        if imagens:

            linhas.append("")

            linhas.append(
                "Detalhes das imagens:"
            )


        for indice_imagem, imagem in enumerate(
            imagens,
            start=1
        ):

            xref = imagem[0]


            try:

                dados = (
                    documento.extract_image(
                        xref
                    )
                )


                extensao = dados.get(
                    "ext",
                    "desconhecida"
                )


                largura = dados.get(
                    "width",
                    0
                )


                altura = dados.get(
                    "height",
                    0
                )


                tamanho = len(
                    dados.get(
                        "image",
                        b""
                    )
                )


                linhas.append(
                    f"  Imagem "
                    f"{indice_imagem:02d}: "
                    f"{extensao.upper()} | "
                    f"{largura}x{altura}px | "
                    f"{formatar_tamanho(tamanho)} | "
                    f"xref={xref}"
                )


            except Exception as erro:

                linhas.append(
                    f"  Imagem "
                    f"{indice_imagem:02d}: "
                    f"erro ao analisar "
                    f"(xref={xref}) — "
                    f"{erro}"
                )


    documento.close()


    linhas.append("")

    linhas.append(
        "=" * 72
    )

    linhas.append(
        "RESUMO"
    )

    linhas.append(
        "=" * 72
    )

    linhas.append(
        f"Total de páginas: "
        f"{total_paginas}"
    )

    linhas.append(
        f"Páginas com texto: "
        f"{paginas_com_texto}"
    )

    linhas.append(
        f"Páginas com imagens: "
        f"{paginas_com_imagens}"
    )

    linhas.append(
        f"Total de imagens incorporadas: "
        f"{total_imagens}"
    )

    linhas.append(
        f"Total de elementos vetoriais: "
        f"{total_desenhos}"
    )

    linhas.append("")

    return "\n".join(
        linhas
    )


# =========================================================
# INVENTÁRIO — DOCX
# =========================================================


def inventariar_docx(
    docx_path: Path
) -> str:
    """
    Analisa um DOCX.

    Registra:
    - quantidade de parágrafos;
    - quantidade de tabelas;
    - quantidade de imagens;
    - nomes das imagens;
    - formatos;
    - tamanhos.
    """

    documento = Document(
        docx_path
    )


    linhas = []


    linhas.append(
        "=" * 72
    )

    linhas.append(
        "VISIUM — INVENTÁRIO DE MATERIAL"
    )

    linhas.append(
        "=" * 72
    )

    linhas.append("")

    linhas.append(
        f"Arquivo: {docx_path.name}"
    )

    linhas.append(
        "Tipo: DOCX"
    )

    linhas.append(
        f"Parágrafos: "
        f"{len(documento.paragraphs)}"
    )

    linhas.append(
        f"Tabelas: "
        f"{len(documento.tables)}"
    )

    linhas.append("")


    # -----------------------------------------------------
    # Imagens incorporadas
    # -----------------------------------------------------

    imagens = []


    with zipfile.ZipFile(
        docx_path,
        "r"
    ) as arquivo_zip:

        for nome in arquivo_zip.namelist():

            if not nome.startswith(
                "word/media/"
            ):

                continue


            if nome.endswith("/"):

                continue


            dados = arquivo_zip.read(
                nome
            )


            imagens.append(
                (
                    nome,
                    dados
                )
            )


    linhas.append(
        "-" * 72
    )

    linhas.append(
        "IMAGENS INCORPORADAS"
    )

    linhas.append(
        "-" * 72
    )

    linhas.append("")


    if not imagens:

        linhas.append(
            "[Nenhuma imagem incorporada encontrada]"
        )


    for indice, (
        nome,
        dados
    ) in enumerate(
        imagens,
        start=1
    ):

        extensao = (
            Path(nome)
            .suffix
            .replace(
                ".",
                ""
            )
            .upper()
        )


        linhas.append(
            f"Imagem {indice:02d}: "
            f"{Path(nome).name} | "
            f"{extensao} | "
            f"{formatar_tamanho(len(dados))}"
        )


    # -----------------------------------------------------
    # Conteúdo textual inicial
    # -----------------------------------------------------

    linhas.append("")

    linhas.append(
        "-" * 72
    )

    linhas.append(
        "CONTEÚDO TEXTUAL"
    )

    linhas.append(
        "-" * 72
    )

    linhas.append("")


    paragrafos_com_texto = [
        paragrafo.text.strip()
        for paragrafo in documento.paragraphs
        if paragrafo.text.strip()
    ]


    if paragrafos_com_texto:

        limite = min(
            10,
            len(paragrafos_com_texto)
        )


        for texto in paragrafos_com_texto[
            :limite
        ]:

            texto_normalizado = (
                texto
                .replace(
                    "\t",
                    " "
                )
                .strip()
            )


            if len(
                texto_normalizado
            ) > 160:

                texto_normalizado = (
                    texto_normalizado[:157]
                    + "..."
                )


            linhas.append(
                f"  - "
                f"{texto_normalizado}"
            )


    else:

        linhas.append(
            "[Nenhum texto extraível encontrado]"
        )


    linhas.append("")

    linhas.append(
        "=" * 72
    )

    linhas.append(
        "RESUMO"
    )

    linhas.append(
        "=" * 72
    )

    linhas.append(
        f"Total de imagens: "
        f"{len(imagens)}"
    )

    linhas.append(
        f"Total de parágrafos: "
        f"{len(documento.paragraphs)}"
    )

    linhas.append(
        f"Total de tabelas: "
        f"{len(documento.tables)}"
    )

    linhas.append("")


    return "\n".join(
        linhas
    )


# =========================================================
# EXTRAÇÃO DE IMAGENS — PDF
# =========================================================


def extrair_imagens_pdf(
    pdf_path: Path
) -> Path:
    """
    Extrai imagens incorporadas de um PDF.
    """

    nome_documento = (
        normalizar_nome(
            pdf_path.stem
        )
    )


    pasta_saida = (
        IMAGES_DIR
        / nome_documento
    )


    pasta_saida.mkdir(
        parents=True,
        exist_ok=True
    )


    documento = fitz.open(
        pdf_path
    )


    total_paginas = len(
        documento
    )

    total_imagens = 0


    print()

    print(
        "=" * 60
    )

    print(
        "VISIUM — EXTRAÇÃO DE IMAGENS"
    )

    print(
        "=" * 60
    )

    print(
        f"Tipo: PDF"
    )

    print(
        f"Arquivo: {pdf_path.name}"
    )

    print(
        f"Páginas: {total_paginas}"
    )

    print(
        f"Destino: {pasta_saida}"
    )

    print(
        "-" * 60
    )


    for numero_pagina in range(
        total_paginas
    ):

        pagina = documento[
            numero_pagina
        ]


        imagens = pagina.get_images(
            full=True
        )


        if not imagens:

            print(
                f"[--] Página "
                f"{numero_pagina + 1:02d}: "
                "nenhuma imagem"
            )

            continue


        for indice_imagem, imagem in enumerate(
            imagens,
            start=1
        ):

            xref = imagem[0]


            try:

                dados = (
                    documento.extract_image(
                        xref
                    )
                )


                extensao = dados[
                    "ext"
                ]


                conteudo = dados[
                    "image"
                ]


                nome_arquivo = (
                    f"pagina-"
                    f"{numero_pagina + 1:02d}"
                    f"-imagem-"
                    f"{indice_imagem:02d}"
                    f".{extensao}"
                )


                caminho_saida = (
                    pasta_saida
                    / nome_arquivo
                )


                if caminho_saida.exists():

                    print(
                        f"[EXISTE] "
                        f"{nome_arquivo}"
                    )

                    continue


                caminho_saida.write_bytes(
                    conteudo
                )


                total_imagens += 1


                largura = dados.get(
                    "width",
                    0
                )


                altura = dados.get(
                    "height",
                    0
                )


                print(
                    f"[OK] "
                    f"{nome_arquivo} "
                    f"({largura}x{altura}px)"
                )


            except Exception as erro:

                print(
                    f"[ERRO] Página "
                    f"{numero_pagina + 1}, "
                    f"imagem "
                    f"{indice_imagem}: "
                    f"{erro}"
                )


    documento.close()


    print(
        "-" * 60
    )

    print(
        f"Finalizado: "
        f"{total_imagens} nova(s) "
        "imagem(ns) extraída(s)."
    )

    print(
        f"Pasta: {pasta_saida}"
    )

    print(
        "=" * 60
    )

    print()


    return pasta_saida


# =========================================================
# EXTRAÇÃO DE IMAGENS — DOCX
# =========================================================


def extrair_imagens_docx(
    docx_path: Path
) -> Path:
    """
    Extrai as imagens originais incorporadas
    em um arquivo DOCX.

    As imagens são armazenadas internamente
    em word/media/.
    """

    nome_documento = (
        normalizar_nome(
            docx_path.stem
        )
    )


    pasta_saida = (
        IMAGES_DIR
        / nome_documento
    )


    pasta_saida.mkdir(
        parents=True,
        exist_ok=True
    )


    total_imagens = 0


    print()

    print(
        "=" * 60
    )

    print(
        "VISIUM — EXTRAÇÃO DE IMAGENS"
    )

    print(
        "=" * 60
    )

    print(
        "Tipo: DOCX"
    )

    print(
        f"Arquivo: {docx_path.name}"
    )

    print(
        f"Destino: {pasta_saida}"
    )

    print(
        "-" * 60
    )


    try:

        with zipfile.ZipFile(
            docx_path,
            "r"
        ) as arquivo_zip:

            arquivos_media = [
                nome
                for nome in arquivo_zip.namelist()
                if (
                    nome.startswith(
                        "word/media/"
                    )
                    and not nome.endswith("/")
                )
            ]


            if not arquivos_media:

                print(
                    "[--] Nenhuma imagem "
                    "incorporada encontrada."
                )


            for indice, nome_interno in enumerate(
                arquivos_media,
                start=1
            ):

                conteudo = (
                    arquivo_zip.read(
                        nome_interno
                    )
                )


                nome_original = (
                    Path(
                        nome_interno
                    ).name
                )


                extensao = (
                    Path(
                        nome_original
                    ).suffix
                    .lower()
                )


                if not extensao:

                    extensao = ".bin"


                nome_arquivo = (
                    f"imagem-"
                    f"{indice:02d}"
                    f"{extensao}"
                )


                caminho_saida = (
                    pasta_saida
                    / nome_arquivo
                )


                if caminho_saida.exists():

                    print(
                        f"[EXISTE] "
                        f"{nome_arquivo}"
                    )

                    continue


                caminho_saida.write_bytes(
                    conteudo
                )


                total_imagens += 1


                print(
                    f"[OK] "
                    f"{nome_arquivo} "
                    f"← "
                    f"{nome_original} "
                    f"({formatar_tamanho(len(conteudo))})"
                )


    except zipfile.BadZipFile as erro:

        print(
            "[ERRO] O arquivo DOCX "
            "não pôde ser lido como ZIP:"
        )

        print(
            erro
        )


        return pasta_saida


    print(
        "-" * 60
    )

    print(
        f"Finalizado: "
        f"{total_imagens} nova(s) "
        "imagem(ns) extraída(s)."
    )

    print(
        f"Pasta: {pasta_saida}"
    )

    print(
        "=" * 60
    )

    print()


    return pasta_saida


# =========================================================
# RELATÓRIO
# =========================================================


def salvar_relatorio(
    documento_path: Path,
    conteudo: str
) -> Path:
    """
    Salva o inventário em um arquivo TXT.
    """

    garantir_pasta_relatorios()


    nome_relatorio = (
        normalizar_nome(
            documento_path.stem
        )
        + "-inventario.txt"
    )


    caminho_relatorio = (
        REPORTS_DIR
        / nome_relatorio
    )


    caminho_relatorio.write_text(
        conteudo,
        encoding="utf-8"
    )


    return caminho_relatorio


# =========================================================
# PROCESSAMENTO — PDF
# =========================================================


def processar_pdf(
    pdf_path: Path
) -> None:

    print()

    print(
        "Analisando material..."
    )

    print(
        pdf_path.name
    )


    relatorio = (
        inventariar_pdf(
            pdf_path
        )
    )


    caminho_relatorio = (
        salvar_relatorio(
            pdf_path,
            relatorio
        )
    )


    print()

    print(
        "[OK] Inventário salvo em:"
    )

    print(
        caminho_relatorio
    )


    extrair_imagens_pdf(
        pdf_path
    )


# =========================================================
# PROCESSAMENTO — DOCX
# =========================================================


def processar_docx(
    docx_path: Path
) -> None:

    print()

    print(
        "Analisando material..."
    )

    print(
        docx_path.name
    )


    relatorio = (
        inventariar_docx(
            docx_path
        )
    )


    caminho_relatorio = (
        salvar_relatorio(
            docx_path,
            relatorio
        )
    )


    print()

    print(
        "[OK] Inventário salvo em:"
    )

    print(
        caminho_relatorio
    )


    extrair_imagens_docx(
        docx_path
    )


# =========================================================
# PROCESSAMENTO COMPLETO
# =========================================================


def processar_documento(
    documento_path: Path
) -> None:

    extensao = (
        documento_path.suffix
        .lower()
    )


    if extensao == ".pdf":

        processar_pdf(
            documento_path
        )

        return


    if extensao == ".docx":

        processar_docx(
            documento_path
        )

        return


    print(
        "[ERRO] Formato não suportado:"
    )

    print(
        documento_path.suffix
    )

    sys.exit(1)


# =========================================================
# EXECUÇÃO
# =========================================================


def main():

    if not DOCS_DIR.exists():

        print(
            "[ERRO] "
            "A pasta docs não existe."
        )

        sys.exit(1)


    documentos = sorted(
        [
            *DOCS_DIR.glob("*.pdf"),
            *DOCS_DIR.glob("*.docx")
        ]
    )


    if not documentos:

        print(
            "[ERRO] "
            "Nenhum PDF ou DOCX "
            "encontrado na pasta docs."
        )

        sys.exit(1)


    # -----------------------------------------------------
    # Documento informado pelo terminal
    # -----------------------------------------------------

    if len(sys.argv) > 1:

        nome_documento = " ".join(
            sys.argv[1:]
        )


        try:

            documento_path = (
                encontrar_documento(
                    nome_documento
                )
            )

        except FileNotFoundError as erro:

            print(
                f"[ERRO] {erro}"
            )

            sys.exit(1)


        processar_documento(
            documento_path
        )

        return


    # -----------------------------------------------------
    # Apenas um documento
    # -----------------------------------------------------

    if len(documentos) == 1:

        processar_documento(
            documentos[0]
        )

        return


    # -----------------------------------------------------
    # Escolha do documento
    # -----------------------------------------------------

    print()

    print(
        "Documentos disponíveis:"
    )

    print()


    for indice, documento in enumerate(
        documentos,
        start=1
    ):

        print(
            f"{indice}. "
            f"{documento.name}"
        )


    print()


    escolha = input(
        "Escolha o documento: "
    ).strip()


    try:

        indice = int(
            escolha
        ) - 1


        documento_path = documentos[
            indice
        ]


    except (
        ValueError,
        IndexError
    ):

        print(
            "[ERRO] "
            "Opção inválida."
        )

        sys.exit(1)


    processar_documento(
        documento_path
    )


# =========================================================
# MAIN
# =========================================================


if __name__ == "__main__":

    main()